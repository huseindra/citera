import pytest
from citera_pipeline.ingest import (
    EmptyDocumentError,
    ExtractionError,
    UnsupportedFileTypeError,
    extract,
)


def test_markdown_passthrough_is_canonical():
    data = "# Title\n\nBody with é and §50.25.\n".encode()
    result = extract("doc.md", data)
    assert result.canonical_text == data.decode()
    assert result.page_map is None


def test_txt_passthrough():
    assert extract("notes.txt", b"plain text").canonical_text == "plain text"


def test_unsupported_type_rejected():
    with pytest.raises(UnsupportedFileTypeError, match="Unsupported file type"):
        extract("consent.rtf", b"whatever")


def test_docx_extracts_with_markdown_style_headings():
    import io

    import docx

    document = docx.Document()
    document.add_heading("What are the risks?", level=2)
    document.add_paragraph("Centraxol has been well tolerated.")
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract("consent.docx", buffer.getvalue())
    assert "## What are the risks?" in result.canonical_text
    assert "Centraxol has been well tolerated." in result.canonical_text
    assert result.page_map is None


def test_corrupt_docx_rejected_cleanly():
    from citera_pipeline.ingest import ExtractionError

    with pytest.raises(ExtractionError, match="Could not read .docx"):
        extract("broken.docx", b"this is not a zip archive")


def test_empty_document_rejected():
    with pytest.raises(EmptyDocumentError):
        extract("empty.md", b"   \n\n  ")


def test_oversized_document_rejected():
    with pytest.raises(ExtractionError, match="20 MB"):
        extract("big.md", b"x" * (20 * 1024 * 1024 + 1))
